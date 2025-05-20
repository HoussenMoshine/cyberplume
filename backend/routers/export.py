# backend/routers/export.py

from fastapi import APIRouter, Depends, HTTPException, status, BackgroundTasks
# MODIFIÉ: Ajout PlainTextResponse, StreamingResponse (pour PDF en mémoire)
from fastapi.responses import FileResponse, PlainTextResponse # StreamingResponse non utilisé finalement
from sqlalchemy.orm import Session, joinedload # Ajout joinedload
import docx # Import python-docx
import os
import tempfile
import html2text # Pour conversion HTML -> Texte simple
# MODIFIÉ: Ajout des imports pour PDF et EPUB
from xhtml2pdf import pisa
from ebooklib import epub
from io import StringIO, BytesIO # Ajout BytesIO pour EPUB en mémoire si besoin
import traceback # Pour l'erreur EPUB
from typing import List # Ajout List

from ..database import get_db
from ..models import Chapter, Project # Ajout Project
from .. import models # Pour accéder aux modèles via models.Project etc.

router = APIRouter(
    prefix="/api",
    tags=["Export"],
    responses={404: {"description": "Not found"}},
)

# --- Helpers ---

# Helper pour convertir HTML en texte simple
def convert_html_to_plain_text(html_content: str) -> str:
    h = html2text.HTML2Text(bodywidth=0) # bodywidth=0 pour éviter les retours à la ligne automatiques
    h.ignore_links = True
    h.ignore_images = True
    h.ignore_emphasis = False # Garder un peu de formatage (gras, italique) si possible
    # Gérer le cas où le contenu est None
    return h.handle(html_content or "")

# Fonction de nettoyage pour BackgroundTask
def remove_file(path: str) -> None:
    try:
        os.remove(path)
        print(f"Temporary file {path} removed.")
    except OSError as e:
        # Ne pas lever d'erreur si le fichier n'existe plus (peut arriver)
        if e.errno != 2: # errno 2: No such file or directory
             print(f"Error removing temporary file {path}: {e}")

# Helper pour générer un nom de fichier sûr
def create_safe_filename(base_title: str, extension: str) -> str:
    """Génère un nom de fichier sûr en remplaçant les caractères non alphanumériques."""
    safe_base = "".join(c if c.isalnum() or c in ('_', '-') else '_' for c in base_title)
    # Limiter la longueur pour éviter les problèmes de système de fichiers
    max_len = 100
    safe_base = safe_base[:max_len]
    # Éviter les noms vides ou juste composés de '_'
    if not safe_base.strip('_'):
        safe_base = f"export_{extension}"
    return f"{safe_base}.{extension}"

# Helper pour récupérer un chapitre ou lever une 404
def get_chapter_or_404(chapter_id: int, db: Session) -> models.Chapter:
    chapter = db.query(models.Chapter).filter(models.Chapter.id == chapter_id).first()
    if not chapter:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Chapter with id {chapter_id} not found",
        )
    return chapter

# NOUVEAU: Helper pour récupérer un projet avec ses chapitres ou lever une 404
def get_project_with_chapters_or_404(project_id: int, db: Session) -> models.Project:
    project = db.query(models.Project).options(
        joinedload(models.Project.chapters) # Charger les chapitres en même temps
    ).filter(models.Project.id == project_id).first()
    if not project:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Project with id {project_id} not found",
        )
    # TODO: Implémenter un tri plus robuste basé sur un champ 'order' dans le futur
    # Pour l'instant, tri par ID comme fallback raisonnable
    project.chapters.sort(key=lambda c: c.id)
    return project

# NOUVEAU: Helper pour combiner le contenu des chapitres en HTML
def combine_chapters_html(project_title: str, chapters: List[models.Chapter]) -> str:
    """Combine les contenus HTML des chapitres en un seul document HTML."""
    combined_html = f"<h1>{project_title}</h1>\n\n"
    for chapter in chapters:
        chapter_title = chapter.title or f"Chapitre {chapter.id}"
        # Utiliser h2 pour les titres de chapitre dans le document combiné
        combined_html += f"<h2>{chapter_title}</h2>\n"
        combined_html += (chapter.content or "<p><i>(Contenu vide)</i></p>") + "\n\n"
    return combined_html

# NOUVEAU: Helper pour générer le contenu XHTML pour EPUB (projet complet)
def create_project_epub_content(project_title: str, chapters: List[models.Chapter]) -> List[epub.EpubHtml]:
    """Crée les objets EpubHtml pour chaque chapitre d'un projet."""
    epub_chapters = []
    # Page titre simple
    title_page_content = f'''<?xml version='1.0' encoding='utf-8'?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="fr" lang="fr">
<head><title>{project_title}</title></head>
<body><h1>{project_title}</h1></body>
</html>'''
    title_page = epub.EpubHtml(title=project_title, file_name='title.xhtml', lang='fr')
    title_page.content = title_page_content.encode('utf-8')
    epub_chapters.append(title_page)

    # Chapitres réels
    for i, chapter in enumerate(chapters):
        chapter_title = chapter.title or f"Chapitre {chapter.id}"
        safe_chapter_filename = f"chap_{chapter.id}.xhtml"
        html_content = chapter.content or "<p><i>(Contenu vide)</i></p>"

        # Contenu XHTML pour chaque chapitre
        epub_content_str = f'''<?xml version='1.0' encoding='utf-8'?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="fr" lang="fr">
<head>
    <meta charset="utf-8"/>
    <title>{chapter_title}</title>
    <style type="text/css">
        body {{ font-family: sans-serif; line-height: 1.5; margin: 1em; }}
        h1, h2 {{ text-align: center; font-size: 1.5em; margin-bottom: 1em; page-break-before: always; }} /* Titre chapitre sur nouvelle page */
        p {{ margin-bottom: 0.5em; text-align: justify; }}
        strong {{ font-weight: bold; }}
        em {{ font-style: italic; }}
    </style>
</head>
<body>
    <h2>{chapter_title}</h2>
    {html_content}
</body>
</html>'''
        epub_chapter = epub.EpubHtml(title=chapter_title, file_name=safe_chapter_filename, lang='fr')
        epub_chapter.content = epub_content_str.encode('utf-8')
        epub_chapters.append(epub_chapter)

    return epub_chapters


# --- Routes d'Export Chapitre (inchangées pour l'instant) ---

@router.get("/chapters/{chapter_id}/export/docx", response_class=FileResponse)
async def export_chapter_docx(chapter_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    """Exporte le contenu d'un chapitre spécifique au format DOCX."""
    chapter = get_chapter_or_404(chapter_id, db)
    doc_title = chapter.title if chapter.title else f"Chapitre_{chapter.id}"

    # Créer un document DOCX en mémoire
    document = docx.Document()
    document.add_heading(doc_title, level=1)

    # Conversion HTML -> Texte simple (temporaire)
    # TODO: Améliorer la conversion HTML -> DOCX pour garder le formatage
    plain_text_content = convert_html_to_plain_text(chapter.content)
    document.add_paragraph(plain_text_content)

    temp_file_path = None
    try:
        fd, temp_file_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        document.save(temp_file_path)

        safe_filename = create_safe_filename(doc_title, "docx")
        headers = {'Content-Disposition': f'attachment; filename="{safe_filename}"'}
        background_tasks.add_task(remove_file, temp_file_path)

        return FileResponse(
            path=temp_file_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=safe_filename,
            headers=headers
        )
    except Exception as e:
        if temp_file_path and os.path.exists(temp_file_path):
            remove_file(temp_file_path) # Utiliser notre helper de suppression
        print(f"Error generating DOCX for chapter {chapter_id}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate DOCX file. Error: {e}",
        )

@router.get("/chapters/{chapter_id}/export/pdf", response_class=FileResponse) # Utilisation de FileResponse pour cohérence
async def export_chapter_pdf(chapter_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    """Exporte le contenu d'un chapitre spécifique au format PDF."""
    chapter = get_chapter_or_404(chapter_id, db)
    doc_title = chapter.title if chapter.title else f"Chapitre_{chapter.id}"
    html_content = chapter.content or ""

    # Ajouter un style de base pour améliorer le rendu PDF
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{doc_title}</title>
        <style>
            @page {{ margin: 1in; }} /* Marges de page */
            body {{ font-family: sans-serif; line-height: 1.5; }}
            h1, h2, h3, h4, h5, h6 {{ font-family: sans-serif; margin-top: 1.5em; margin-bottom: 0.5em; page-break-after: avoid; }}
            h1 {{ font-size: 24pt; text-align: center; margin-bottom: 1em; }}
            p {{ margin-bottom: 1em; text-align: justify; }}
            strong, b {{ font-weight: bold; }}
            em, i {{ font-style: italic; }}
            /* Ajoutez d'autres styles si nécessaire */
        </style>
    </head>
    <body>
        <h1>{doc_title}</h1>
        {html_content}
    </body>
    </html>
    """

    temp_file_path = None
    try:
        # Créer un fichier temporaire pour le PDF
        fd, temp_file_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)

        # Ouvrir le fichier temporaire en mode binaire pour l'écriture
        with open(temp_file_path, "wb") as pdf_file:
            # Convertir HTML en PDF
            pisa_status = pisa.CreatePDF(
                StringIO(styled_html), # Source HTML (StringIO pour gérer la string)
                dest=pdf_file,         # Destination (fichier binaire)
                encoding='utf-8'       # Encodage important
            )

        if pisa_status.err:
            # Supprimer le fichier temporaire en cas d'erreur de génération PDF
            remove_file(temp_file_path)
            raise Exception(f"PDF generation error code: {pisa_status.err}")

        safe_filename = create_safe_filename(doc_title, "pdf")
        headers = {'Content-Disposition': f'attachment; filename="{safe_filename}"'}
        background_tasks.add_task(remove_file, temp_file_path)

        return FileResponse(
            path=temp_file_path,
            media_type='application/pdf',
            filename=safe_filename,
            headers=headers
        )

    except Exception as e:
        # Assurer la suppression si le fichier existe encore
        if temp_file_path and os.path.exists(temp_file_path):
            remove_file(temp_file_path)
        print(f"Error generating PDF for chapter {chapter_id}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate PDF file. Error: {e}",
        )


@router.get("/chapters/{chapter_id}/export/txt", response_class=PlainTextResponse)
async def export_chapter_txt(chapter_id: int, db: Session = Depends(get_db)):
    """Exporte le contenu d'un chapitre spécifique au format TXT."""
    chapter = get_chapter_or_404(chapter_id, db)
    doc_title = chapter.title if chapter.title else f"Chapitre_{chapter.id}"

    try:
        plain_text_content = convert_html_to_plain_text(chapter.content)
        # Ajouter le titre au début du texte
        full_text = f"{doc_title}\n\n{plain_text_content}"

        safe_filename = create_safe_filename(doc_title, "txt")
        headers = {'Content-Disposition': f'attachment; filename="{safe_filename}"'}

        # Retourner directement le texte avec les bons headers
        return PlainTextResponse(content=full_text, headers=headers, media_type='text/plain; charset=utf-8') # Spécifier UTF-8

    except Exception as e:
        print(f"Error generating TXT for chapter {chapter_id}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate TXT file. Error: {e}",
        )

@router.get("/chapters/{chapter_id}/export/epub", response_class=FileResponse)
async def export_chapter_epub(chapter_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    """Exporte le contenu d'un chapitre spécifique au format EPUB."""
    chapter = get_chapter_or_404(chapter_id, db)
    doc_title = chapter.title if chapter.title else f"Chapitre_{chapter.id}"
    html_content = chapter.content or ""

    temp_file_path = None
    try:
        # Créer le livre EPUB
        book = epub.EpubBook()

        # Définir les métadonnées
        book.set_identifier(f"cyberplume-chapter-{chapter_id}-{tempfile.gettempprefix()}") # ID plus unique
        book.set_title(doc_title)
        book.set_language('fr') # Langue par défaut, pourrait être configurable
        book.add_author("CyberPlume User") # Auteur par défaut

        # Créer le chapitre EPUB à partir du HTML
        safe_chapter_filename = f"chap_{chapter_id}.xhtml"
        epub_chapter = epub.EpubHtml(title=doc_title, file_name=safe_chapter_filename, lang='fr')

        # MODIFIÉ: Création du contenu XHTML pour EPUB
        epub_content_str = f'''<?xml version='1.0' encoding='utf-8'?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="fr" lang="fr">
<head>
    <meta charset="utf-8"/>
    <title>{doc_title}</title>
    <style type="text/css">
        /* Styles de base pour EPUB */
        body {{ font-family: sans-serif; line-height: 1.5; margin: 1em; }}
        h1 {{ text-align: center; font-size: 1.5em; margin-bottom: 1em;}}
        p {{ margin-bottom: 0.5em; text-align: justify; }}
        strong {{ font-weight: bold; }}
        em {{ font-style: italic; }}
        /* Ajoutez d'autres styles si nécessaire */
    </style>
</head>
<body>
    <h1>{doc_title}</h1>
    {html_content}
</body>
</html>'''
        # AJOUT: Log du contenu avant encodage pour débogage
        # print(f"--- DEBUG EPUB HTML Content for chapter {chapter_id} ---")
        # print(epub_content_str)
        # print("--- END DEBUG EPUB HTML Content ---")
        epub_chapter.content = epub_content_str.encode('utf-8') # Encoder explicitement en UTF-8 bytes

        # Ajouter le chapitre au livre
        book.add_item(epub_chapter)

        # Définir la table des matières et le 'spine' (ordre de lecture)
        book.toc = (epub.Link(safe_chapter_filename, doc_title, f'chap-{chapter_id}'),)
        book.spine = ['nav', epub_chapter] # 'nav' est la table des matières générée automatiquement

        # Ajouter la page de navigation NCX et Cover (optionnel)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Créer un fichier temporaire pour l'EPUB
        fd, temp_file_path = tempfile.mkstemp(suffix=".epub")
        os.close(fd)

        # Écrire le livre EPUB dans le fichier temporaire
        epub.write_epub(temp_file_path, book, {})

        safe_filename = create_safe_filename(doc_title, "epub")
        headers = {'Content-Disposition': f'attachment; filename="{safe_filename}"'}
        background_tasks.add_task(remove_file, temp_file_path)

        return FileResponse(
            path=temp_file_path,
            media_type='application/epub+zip',
            filename=safe_filename,
            headers=headers
        )

    except Exception as e:
        if temp_file_path and os.path.exists(temp_file_path):
            remove_file(temp_file_path)
        print(f"Error generating EPUB for chapter {chapter_id}: {e}")
        # Afficher plus de détails sur l'erreur pour le débogage
        traceback.print_exc()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate EPUB file. Error: {e}",
        )

# --- NOUVEAU: Routes d'Export Projet Complet ---

@router.get("/projects/{project_id}/export/docx", response_class=FileResponse)
async def export_project_docx(project_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    """Exporte le contenu complet d'un projet au format DOCX."""
    project = get_project_with_chapters_or_404(project_id, db)
    doc_title = project.title if project.title else f"Projet_{project.id}"

    # Créer un document DOCX
    document = docx.Document()
    document.add_heading(doc_title, level=0) # Titre du projet

    # Ajouter chaque chapitre
    for chapter in project.chapters:
        chapter_title = chapter.title or f"Chapitre {chapter.id}"
        document.add_heading(chapter_title, level=1) # Titre du chapitre
        # TODO: Améliorer la conversion HTML -> DOCX pour garder le formatage
        plain_text_content = convert_html_to_plain_text(chapter.content)
        document.add_paragraph(plain_text_content)
        document.add_page_break() # Saut de page entre les chapitres

    temp_file_path = None
    try:
        fd, temp_file_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        document.save(temp_file_path)

        safe_filename = create_safe_filename(doc_title, "docx")
        headers = {'Content-Disposition': f'attachment; filename="{safe_filename}"'}
        background_tasks.add_task(remove_file, temp_file_path)

        return FileResponse(
            path=temp_file_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=safe_filename,
            headers=headers
        )
    except Exception as e:
        if temp_file_path and os.path.exists(temp_file_path):
            remove_file(temp_file_path)
        print(f"Error generating DOCX for project {project_id}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate project DOCX file. Error: {e}",
        )

@router.get("/projects/{project_id}/export/pdf", response_class=FileResponse)
async def export_project_pdf(project_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    """Exporte le contenu complet d'un projet au format PDF."""
    project = get_project_with_chapters_or_404(project_id, db)
    doc_title = project.title if project.title else f"Projet_{project.id}"

    # Combiner le HTML des chapitres
    combined_html_content = ""
    for chapter in project.chapters:
        chapter_title = chapter.title or f"Chapitre {chapter.id}"
        combined_html_content += f"<h2>{chapter_title}</h2>\n" # Utiliser h2 pour les titres de chapitre
        combined_html_content += (chapter.content or "<p><i>(Contenu vide)</i></p>") + "\n<hr />\n" # Séparateur simple

    # Ajouter un style de base pour améliorer le rendu PDF
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{doc_title}</title>
        <style>
            @page {{ margin: 1in; }}
            body {{ font-family: sans-serif; line-height: 1.5; }}
            h1 {{ font-size: 24pt; text-align: center; margin-bottom: 1.5em; page-break-after: avoid; }} /* Titre projet */
            h2 {{ font-size: 18pt; margin-top: 2em; margin-bottom: 1em; page-break-before: always; page-break-after: avoid; }} /* Titre chapitre */
            p {{ margin-bottom: 1em; text-align: justify; }}
            strong, b {{ font-weight: bold; }}
            em, i {{ font-style: italic; }}
            hr {{ page-break-after: always; border: 0; }} /* Saut de page après chaque chapitre via hr */
        </style>
    </head>
    <body>
        <h1>{doc_title}</h1>
        {combined_html_content}
    </body>
    </html>
    """

    temp_file_path = None
    try:
        fd, temp_file_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        with open(temp_file_path, "wb") as pdf_file:
            pisa_status = pisa.CreatePDF(StringIO(styled_html), dest=pdf_file, encoding='utf-8')

        if pisa_status.err:
            remove_file(temp_file_path)
            raise Exception(f"PDF generation error code: {pisa_status.err}")

        safe_filename = create_safe_filename(doc_title, "pdf")
        headers = {'Content-Disposition': f'attachment; filename="{safe_filename}"'}
        background_tasks.add_task(remove_file, temp_file_path)

        return FileResponse(
            path=temp_file_path,
            media_type='application/pdf',
            filename=safe_filename,
            headers=headers
        )
    except Exception as e:
        if temp_file_path and os.path.exists(temp_file_path):
            remove_file(temp_file_path)
        print(f"Error generating PDF for project {project_id}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate project PDF file. Error: {e}",
        )

@router.get("/projects/{project_id}/export/txt", response_class=PlainTextResponse)
async def export_project_txt(project_id: int, db: Session = Depends(get_db)):
    """Exporte le contenu complet d'un projet au format TXT."""
    project = get_project_with_chapters_or_404(project_id, db)
    doc_title = project.title if project.title else f"Projet_{project.id}"

    try:
        full_text = f"{doc_title}\n\n"
        full_text += ("=" * len(doc_title)) + "\n\n" # Séparateur titre

        for chapter in project.chapters:
            chapter_title = chapter.title or f"Chapitre {chapter.id}"
            plain_text_content = convert_html_to_plain_text(chapter.content)
            full_text += f"{chapter_title}\n"
            full_text += ("-" * len(chapter_title)) + "\n" # Séparateur chapitre
            full_text += f"{plain_text_content}\n\n"

        safe_filename = create_safe_filename(doc_title, "txt")
        headers = {'Content-Disposition': f'attachment; filename="{safe_filename}"'}

        return PlainTextResponse(content=full_text, headers=headers, media_type='text/plain; charset=utf-8')
    except Exception as e:
        print(f"Error generating TXT for project {project_id}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate project TXT file. Error: {e}",
        )

@router.get("/projects/{project_id}/export/epub", response_class=FileResponse)
async def export_project_epub(project_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    """Exporte le contenu complet d'un projet au format EPUB."""
    project = get_project_with_chapters_or_404(project_id, db)
    doc_title = project.title if project.title else f"Projet_{project.id}"

    temp_file_path = None
    try:
        # Créer le livre EPUB
        book = epub.EpubBook()
        book.set_identifier(f"cyberplume-project-{project_id}-{tempfile.gettempprefix()}")
        book.set_title(doc_title)
        book.set_language('fr')
        book.add_author("CyberPlume User")

        # Créer les chapitres EPUB
        epub_items = create_project_epub_content(doc_title, project.chapters)

        # Ajouter les items au livre
        for item in epub_items:
            book.add_item(item)

        # Définir la table des matières et le 'spine'
        # La TOC inclut la page titre et les chapitres
        book.toc = [epub.Link(item.file_name, item.title, item.file_name.split('.')[0]) for item in epub_items]
        # Le spine définit l'ordre de lecture
        book.spine = ['nav'] + epub_items # 'nav' est la TOC générée

        # Ajouter la page de navigation NCX et Nav
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Créer un fichier temporaire
        fd, temp_file_path = tempfile.mkstemp(suffix=".epub")
        os.close(fd)

        # Écrire l'EPUB
        epub.write_epub(temp_file_path, book, {})

        safe_filename = create_safe_filename(doc_title, "epub")
        headers = {'Content-Disposition': f'attachment; filename="{safe_filename}"'}
        background_tasks.add_task(remove_file, temp_file_path)

        return FileResponse(
            path=temp_file_path,
            media_type='application/epub+zip',
            filename=safe_filename,
            headers=headers
        )
    except Exception as e:
        if temp_file_path and os.path.exists(temp_file_path):
            remove_file(temp_file_path)
        print(f"Error generating EPUB for project {project_id}: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate project EPUB file. Error: {e}",
        )